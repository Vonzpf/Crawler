#coding:utf-8

import urllib
import re
#import os
from bs4 import BeautifulSoup
#from bs4.builder._htmlparser import BeautifulSoupHTMLParser
from docx import Document
#from docx.shared import Inches

def getHtml(url):
    page = urllib.urlopen(url)
    html = page.read()
    return html

def getanswer(html):
    reg = r'data-entry-url=".+"'
    urlre = re.compile(reg)
    urllist = re.findall(urlre,html)
    x = 0
    title_doc = Document()
    for answerurl in urllist:
        urlsplited = answerurl.split('"')
        fullurl = 'https://www.zhihu.com' + urlsplited[1]   #获取完整的答案URL
        
        answerpage = urllib.urlopen(fullurl)
        answerhtml = answerpage.read()
        
        soup = BeautifulSoup(answerhtml,"lxml")
        title_tag = soup.find("h2", class_="zm-item-title")
        answer_tag = soup.find("div", class_="zm-editable-content clearfix")
        
        title_tag.a.unwrap()
        
        #处理回答中不必要的标签
        while(answer_tag.find_all("noscript")):
            answer_tag.noscript.extract()
        while(answer_tag.find_all("br")):
            answer_tag.br.decompose()
        while(answer_tag.find_all("b")):
            answer_tag.b.unwrap()
        while(answer_tag.find_all("u")):
            answer_tag.u.unwrap()
        while(answer_tag.find_all("p")):
            answer_tag.p.unwrap()
        
        title = unicode(x)+title_tag.get_text()
        print(title)
        
        answer = answer_tag.prettify()
        
        #保存为word
        document = Document()
        document.add_paragraph(answer)
        document.save('d:\%s.docx' %x)
        
        title_doc.add_paragraph(title)
        x += 1

    title_doc.save('d:\\alltitle.docx')

html = getHtml("https://www.zhihu.com/collection/47588314?page=1")

getanswer(html)